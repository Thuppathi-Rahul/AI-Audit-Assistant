import pandas as pd
from pydantic import BaseModel, Field
from typing import Literal
import requests
import datetime
from langchain_openai import ChatOpenAI
from langchain.tools import StructuredTool
from langchain.agents import AgentExecutor, create_openai_tools_agent
from langchain import hub
from PyPDF2 import PdfReader
from docx import Document
import streamlit as st
import io
import os
import plotly.graph_objects as go
from openpyxl.styles import Alignment
from shareplum import Site, Office365
from github import Github, Auth 


# --- STATIC DATA (The Single Source of Truth - Corrected with all keywords) ---
AUDIT_CHECKLIST = [
    {"subject": "Project Initiation", "question": "Is the Signed SOW and MSA available also verify the change orders if any?", "keywords": ["sow", "msa", "checklist"], "weight": 3, "tags": ["PCI", "GDPR", "CMMI", "ITSM"]},
    {"subject": "Inception and Discovery", "question": "Is the High Level Architecture understood and documented?", "keywords": ["high level design", "hld", "architecture"], "weight": 2, "tags": ["PCI", "CMMI"]},
    {"subject": "Inception and Discovery", "question": "Is there a high level release plan available including high level Estimates?", "keywords": ["agile estimation", "release planning", "estimates"], "weight": 1, "tags": ["PCI", "CMMI"]},
    {"subject": "Inception and Discovery", "question": "Are the non-functional requirements identified?", "keywords": ["jira", "product backlog", "user stories", "nfr"], "weight": 3, "tags": ["PCI", "CMMI"]},
    {"subject": "Inception and Discovery", "question": "Is the Project process (with required tailoring) that need to be followed are identified?", "keywords": ["pmp", "project management plan"], "weight": 2, "tags": ["PCI", "CMMI"]},
    {"subject": "Sprint 0", "question": "Is the Project Management Plan and Quality Plan defined for this project?", "keywords": ["pmp", "project management plan", "gdq-qa", "quality plan"], "weight": 1, "tags": ["PCI", "GDPR", "CMMI", "ITSM"]},
    {"subject": "Sprint 0", "question": "Is the change management planning, customer supplied assets, NDA and Information Security related aspects defined?", "keywords": ["project management process", "change management", "nda"], "weight": 3, "tags": ["PCI", "GDPR", "CMMI", "ITSM"]},
    {"subject": "Sprint 0", "question": "Does the PMP have Risk Management and Issue Resolution plans?", "keywords": ["risk register", "pmp", "project management plan"], "weight": 2, "tags": ["PCI", "GDPR", "CMMI", "ITSM"]},
    {"subject": "Sprint 0", "question": "Does the Quality Plan have the 1. Audits and Review plan defined 2. Measurement plan / agile metrics goals defined 3. Phase Gates planned (PG7 [Design Completion Review], PG8 [Production/Go-Live Readiness] and PG9[Project Closure])", "keywords": ["gdq-qa", "plan", "phase gate", "pg7", "pg8", "pg9", "score card"], "weight": 1, "tags": ["PCI", "CMMI", "ITSM"]},
    {"subject": "Sprint 0", "question": "Is the PMP (Project Management Plan) reviewed and approved by the Service Line Manager and QA team?", "keywords": ["pmp", "project management plan"], "weight": 3, "tags": ["PCI", "GDPR", "CMMI", "ITSM"]},
    {"subject": "Sprint 0", "question": "Is Definition of Done define?", "keywords": ["definition of done", "dod"], "weight": 2, "tags": ["PCI", "CMMI"]},
    {"subject": "Sprint 0", "question": "Did team start developing the user stories? Are the user stories elaborate, clear to estimate?", "keywords": ["user stories", "design", "develop"], "weight": 1, "tags": ["PCI", "CMMI"]},
    {"subject": "Sprint 0", "question": "Is the acceptance criteria defined for User stories?", "keywords": ["user stories", "acceptance criteria"], "weight": 3, "tags": ["PCI", "CMMI"]},
    {"subject": "Sprint 0", "question": "Are user stories and acceptance criteria reviewed and approved by product owner?", "keywords": ["user stories", "acceptance criteria", "jira"], "weight": 2, "tags": ["PCI", "CMMI"]},
    {"subject": "Sprint 0", "question": "Are the Product owner, Scrum master and Scrum team identified for the project?", "keywords": ["working agreement", "roles", "responsibilities"], "weight": 1, "tags": ["PCI", "CMMI"]},
    {"subject": "Sprint Planning", "question": "Did the team estimate for user stories, in terms of story points and efforts? Did the team estimate to granular level?", "keywords": ["agile estimation", "release planning", "story points"], "weight": 3, "tags": ["PCI", "CMMI"]},
    {"subject": "Sprint Planning", "question": "Was the entire team involved in estimation activities, including Product owner, Scrum master and Scrum Team?", "keywords": ["agile estimation", "sprint planning"], "weight": 2, "tags": ["PCI", "CMMI"]},
    {"subject": "Sprint Planning", "question": "Are the user stories Reviewed and approved by Product owner?", "keywords": ["sprint backlog", "burndown", "user stories"], "weight": 1, "tags": ["PCI", "GDPR", "CMMI"]},
    {"subject": "Sprint Execution", "question": "Did the team prepare low level design for all the functional user stories?", "keywords": ["detail design", "lld"], "weight": 3, "tags": ["PCI", "CMMI"]},
    {"subject": "Sprint Execution", "question": "Is LLD reviewed and approved by SME/ Product owner", "keywords": ["tca.020", "technical architecture", "high level design"], "weight": 2, "tags": ["PCI", "CMMI"]},
    {"subject": "Sprint Execution", "question": "Did team perform unit testing of the developed code?", "keywords": ["unit test plan"], "weight": 1, "tags": ["PCI", "CMMI"]},
    {"subject": "Sprint Execution", "question": "Did product owner reviewed and approved the Test cases?", "keywords": ["unit test plan", "test cases"], "weight": 3, "tags": ["PCI", "CMMI"]},
    {"subject": "Sprint Execution", "question": "Is daily standup meeting planned and conducted?", "keywords": ["daily standup", "meeting template"], "weight": 2, "tags": ["PCI", "CMMI"]},
    {"subject": "Sprint Execution", "question": "Are the relevant stakeholders, Product owner, Scrum master and team part of the standup meeting?", "keywords": ["daily standup", "impediments list"], "weight": 1, "tags": ["PCI", "CMMI"]},
    {"subject": "Sprint Execution", "question": "Are the CI & Non CI needs identified and implemented?", "keywords": ["pmp", "project management plan", "ci"], "weight": 3, "tags": ["PCI", "CMMI"]},
    {"subject": "Project Status Reporting/ PG6", "question": "Is project status reviewed with senior management at appropriate intervals? a. Overall status b. Project performance (achievements & milestones) c. Open issues d. Risks e. Action items f. Cost & time performance against plan g. Quality metrics i. Team member's skill assessment report j. IQA and CQA results", "keywords": ["risk register", "rail", "rolling action", "phase gate 6", "hi-dash"], "weight": 2, "tags": ["PCI", "CMMI"]},
    {"subject": "Qualitative Assurance", "question": "Are the metrics captured and reported for each Sprint?", "keywords": ["agile metrics", "evm", "hi-dash"], "weight": 1, "tags": ["PCI", "CMMI"]},
    {"subject": "Risk Management", "question": "Are all risks identified and documented?", "keywords": ["risk register", "rail", "phase gate 6"], "weight": 3, "tags": ["PCI", "CMMI"]},
    {"subject": "Risk Management", "question": "Are Mitigation and Contingency Plans in place?", "keywords": ["risk register", "mitigation", "contingency"], "weight": 2, "tags": ["PCI", "CMMI"]},
    {"subject": "Risk Management", "question": "Are risks reviewed and updated periodically.?", "keywords": ["risk register"], "weight": 1, "tags": ["PCI", "CMMI"]},
    {"subject": "Risk Management", "question": "Are Mitigation plans effective. If risks had occurred, look for the implementation of contingency plan for critical risks and impact assessment ?", "keywords": ["risk register", "mitigation", "contingency"], "weight": 3, "tags": ["PCI", "CMMI"]},
    {"subject": "Customer Complaints & CSS", "question": "Is the Progress on action plan tracked periodically and the associated risk also updated?", "keywords": ["project status report", "hi-dash"], "weight": 2, "tags": ["PCI", "CMMI"]},
    {"subject": "Customer Complaints & CSS", "question": "Has there been a CSS initiated for the project in the last 6 months?", "keywords": ["css", "email communication"], "weight": 1, "tags": ["PCI", "CMMI"]},
    {"subject": "Phase Gate and Code Quality Compliance", "question": "Are Code Quality Audits planned and conducted for this project as per frequency defined in PMP?", "keywords": ["pmp", "code quality", "cqa", "checklist", "rail", "irf tool"], "weight": 3, "tags": ["PCI", "CMMI"]},
    {"subject": "Phase Gate and Code Quality Compliance", "question": "Has the PG7 (Design Completion Review) been conducted as planned and action items tracked to closure", "keywords": ["phase gate 7", "pg7", "rail"], "weight": 2, "tags": ["PCI", "CMMI"]},
    {"subject": "Phase Gate and Code Quality Compliance", "question": "Has the PG8 (Production/ Go-Live Readiness) been conducted as planned and action items tracked to closure", "keywords": ["scorecard", "pg8", "hi-dash"], "weight": 1, "tags": ["PCI", "CMMI"]},
    {"subject": "Phase Gate and Code Quality Compliance", "question": "Has the PG9 (Project Closure) been conducted as planned and lessons learned/ key success factors documented", "keywords": ["pg9", "project closure", "report"], "weight": 3, "tags": ["PCI", "CMMI", "ITSM"]},
    {"subject": "Information Security (Bare Minimum Checks)", "question": "Are Information Security related needs, client expectations, requirements identified in Project Management Plan?", "keywords": ["pmp", "project management plan", "information security"], "weight": 2, "tags": ["PCI", "GDPR", "Infosec"]},
    {"subject": "Information Security", "question": "Is Project Team aware of Information Security related policies like Clean/Clear Desk, Password Management etc.? Did they attend ISMS Training Sessions Conducted by Infosec team?", "keywords": ["global information security policy"], "weight": 1, "tags": ["PCI", "GDPR", "Infosec"]},
    {"subject": "Information Security (Bare Minimum Checks)", "question": "Are Information Security Risks identified and monitored to closure with Proper Mitigation Plans as per CIA?", "keywords": ["risk register", "information security"], "weight": 3, "tags": ["PCI", "GDPR", "Infosec"]},
    {"subject": "Information Security (Bare Minimum Checks)", "question": "Are Information Security Audits conducted as per defined frequency in PMP ( As Applicable)?", "keywords": ["pmp", "project management plan", "information security audit"], "weight": 2, "tags": ["PCI", "GDPR", "Infosec"]},
    {"subject": "Information Security (Bare Minimum Checks)", "question": "Is the project's purpose and setup clearly documented in the README.md file?", "keywords": ["README.md"], "weight": 2, "tags": ["PCI", "Infosec", "GitHub"], "source": "github"},
    {"subject": "Information Security (Bare Minimum Checks)", "question": "Does the database connection file contain any hardcoded passwords or secrets?", "keywords": ["config.py", "settings.py", "db.py"], "weight": 3, "tags": ["PCI", "Infosec", "GitHub"], "source": "github"}

]



# --- DOCUMENT EXTRACTION ---
def extract_text_from_pdf(file_bytes):
    try:
        pdf_reader = PdfReader(io.BytesIO(file_bytes))
        text = "".join(page.extract_text() or "" for page in pdf_reader.pages)
        return text
    except Exception as e:
        return f"Error reading PDF: {e}"

def extract_text_from_docx(file_bytes):
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join(para.text for para in doc.paragraphs)
        return text
    except Exception as e:
        return f"Error reading DOCX: {e}"
# --- SharePoint Document Fetching ---
def fetch_sharepoint_docs(site_url, folder_path):
    """Connects to SharePoint using credentials from .env file."""
    extracted_texts = {}
    
    # --- CHANGE: Read credentials from environment variables ---
    username = os.getenv("SHAREPOINT_USERNAME")
    password = os.getenv("SHAREPOINT_PASSWORD")

    if not username or not password:
        st.error("SharePoint credentials not found in .env file. Please add SHAREPOINT_USERNAME and SHAREPOINT_PASSWORD.")
        return {}

    try:
        authcookie = Office365(site_url, username=username, password=password).GetCookies()
        site = Site(site_url, authcookie=authcookie)
        folder = site.Folder(folder_path)
        files = folder.files

        for file_info in files:
            file_name = file_info['Name']
            if file_name.lower().endswith(('.pdf', '.docx')):
                st.write(f"-> Found '{file_name}' in SharePoint. Downloading...")
                file_content = folder.get_file(file_name)
                
                if file_name.lower().endswith('.pdf'):
                    extracted_texts[file_name] = extract_text_from_pdf(file_content)
                elif file_name.lower().endswith('.docx'):
                    extracted_texts[file_name] = extract_text_from_docx(file_content)
        
        return extracted_texts

    except Exception as e:
        st.error(f"Failed to connect or download from SharePoint: {e}")
        return {}
# --- GitHub File Fetching ---
def fetch_github_file_content(repo_name: str, file_paths: list) -> str:
    """Connects to GitHub and reads the content of specific files."""
    token = os.getenv("GITHUB_TOKEN")
    if not token:
        st.error("GitHub token not found in .env file. Please add GITHUB_TOKEN.")
        return "Error: GitHub token not configured."
    
    all_content = []
    try:
        auth = Auth.Token(token)
        g = Github(auth=auth)
        repo = g.get_repo(repo_name)
        
        for file_path in file_paths:
            try:
                # This inner try handles cases where one of the files doesn't exist
                content_file = repo.get_contents(file_path)
                decoded_content = content_file.decoded_content.decode("utf-8")
                all_content.append(f"--- Content from {file_path} ---\n{decoded_content}")
            except Exception as file_error:
                # If a specific file is not found, we note it but don't crash.
                print(f"Could not find file '{file_path}' in repo '{repo_name}': {file_error}")
                all_content.append(f"--- Content from {file_path} ---\nError: The file '{file_path}' was not found in the repository.")
        
        return "\n\n".join(all_content)

    except Exception as e:
        # This outer try handles bigger errors like a bad repo name or bad credentials
        st.error(f"Failed to connect to GitHub repo '{repo_name}': {e}")
        return f"Error: Could not connect to GitHub repository {repo_name}."

# --- API COMMUNICATION ---
BACKEND_URL = "http://127.0.0.1:8000"

def update_irf_and_ui(question: str, answer: str, explanation: str) -> str:
    timestamp = datetime.datetime.now(datetime.timezone.utc)
    run_id = st.session_state.get("run_id", "default_run")
    payload = { "run_id": run_id, "question": question, "answer": answer, "explanation": explanation, "timestamp": timestamp.isoformat() }
    try:
        response = requests.post(f"{BACKEND_URL}/submit_finding/", json=payload)
        response.raise_for_status()
        new_row = pd.DataFrame([{"Question": question, "Answer": answer, "Explanation": explanation}])
        if "audit_results" not in st.session_state:
            st.session_state.audit_results = pd.DataFrame(columns=["Question", "Answer", "Explanation"])
        st.session_state.audit_results = pd.concat([st.session_state.audit_results, new_row], ignore_index=True)
        return f"Successfully submitted finding to IRF tool. Response: {response.json()}"
    except requests.exceptions.RequestException as e:
        return f"Failed to submit finding to IRF tool. Error: {e}"

# --- AGENT SETUP ---
class AuditFindingInput(BaseModel):
    question: str = Field(description="The full text of the audit question that was answered.")
    answer: Literal["Yes", "No", "Partial"] = Field(description="The final answer based on the context.")
    explanation: str = Field(description="A short explanation justifying the answer.")
# --- Simple LLM instance for direct calls without agent logic ---
@st.cache_resource
def get_llm():
    print("INFO: Creating new ChatOpenAI instance.")
    return ChatOpenAI(model_name="gpt-4-turbo", temperature=0)

@st.cache_resource
def get_agent_executor():
    print("INFO: Creating new LangChain AgentExecutor instance.")
    llm = ChatOpenAI(model_name="gpt-4-turbo", temperature=0)
    tools = [ StructuredTool.from_function( func=update_irf_and_ui, name="SubmitAuditFinding", description="Use this tool to submit the final answer for a single audit question.", args_schema=AuditFindingInput ) ]
    agent_prompt = hub.pull("hwchase17/openai-tools-agent")
    agent = create_openai_tools_agent(llm, tools, agent_prompt)
    agent_executor = AgentExecutor(agent=agent, tools=tools, verbose=True)
    return agent_executor

# --- SCORING, CHARTING, AND EXCEL FUNCTIONS ---
def to_excel(data: list) -> bytes:
    if not data:
        return b""
    answer_multipliers = {"Yes": 1.0, "Partial": 0.5, "No": 0.0}
    current_checklist = AUDIT_CHECKLIST + st.session_state.get("custom_checklist", [])
    question_to_weight_map = {item['question']: item.get('weight', 0) for item in current_checklist}
    
    report_data = []
    for finding in data:
        weight = question_to_weight_map.get(finding['question'], 0)
        multiplier = answer_multipliers.get(finding.get('answer'), 0.0)
        score = weight * multiplier
        new_finding = finding.copy()
        new_finding['weight'] = weight
        new_finding['score'] = score
        report_data.append(new_finding)
        
    df = pd.DataFrame(report_data)
    df_report = df[['question', 'weight', 'answer', 'score', 'explanation', 'timestamp']].copy()
    df_report.rename(columns={'question': 'Audit Question','weight': 'Weightage','answer': 'Answer','score': 'Achieved Score','explanation': 'Explanation/Comments','timestamp': 'Timestamp (UTC)'}, inplace=True)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_report.to_excel(writer, index=False, sheet_name='Audit_Results')
        worksheet = writer.sheets['Audit_Results']
        worksheet.column_dimensions['A'].width = 80
        worksheet.column_dimensions['B'].width = 12
        worksheet.column_dimensions['C'].width = 15
        worksheet.column_dimensions['D'].width = 18
        worksheet.column_dimensions['E'].width = 80
        worksheet.column_dimensions['F'].width = 20
        wrap_alignment = Alignment(wrap_text=True, vertical='top')
        for row in worksheet.iter_rows(min_row=2, max_col=6): 
            for cell in [row[0], row[4]]:
                cell.alignment = wrap_alignment
    return output.getvalue()

def calculate_all_scores(run_data, selected_run_id):
    scores = {}
    answer_multipliers = {"Yes": 1.0, "Partial": 0.5, "No": 0.0}
    
    current_checklist = AUDIT_CHECKLIST + st.session_state.get("custom_checklist", [])
    compliance_areas = ["PCI", "GDPR", "Infosec", "CMMI", "ITSM", "Custom"]
    
    for area in compliance_areas:
        total_score_achieved = 0.0
        max_possible_score = 0.0
        
        area_checklist = [item for item in current_checklist if area in item.get("tags", [])]
        question_to_idx_map = {item['question']: i for i, item in enumerate(current_checklist)}
        
        for item in area_checklist:
            weight = item.get('weight', 0)
            question_idx = question_to_idx_map.get(item['question'], -1)
            
            current_answer = st.session_state.get(f"answer_{question_idx}_{selected_run_id}")
            if current_answer is None:
                finding = run_data.get(item['question'])
                if finding: current_answer = finding.get('answer')

            if current_answer != "N/A":
                max_possible_score += weight
                multiplier = answer_multipliers.get(current_answer, 0.0)
                total_score_achieved += weight * multiplier

        if max_possible_score > 0:
            percentage = (total_score_achieved / max_possible_score) * 100
        else:
            percentage = 0.0

        scores[area] = { "percentage": percentage, "achieved": total_score_achieved, "max": max_possible_score }
        
    return scores

def get_answer_counts(run_data, selected_run_id, compliance_area):
    counts = {"Yes": 0, "No": 0, "Partial": 0}
    current_checklist = AUDIT_CHECKLIST + st.session_state.get("custom_checklist", [])
    area_checklist = [item for item in current_checklist if compliance_area in item.get("tags", [])]
    question_to_idx_map = {item['question']: i for i, item in enumerate(current_checklist)}

    for item in area_checklist:
        question_idx = question_to_idx_map.get(item['question'], -1)
        current_answer = st.session_state.get(f"answer_{question_idx}_{selected_run_id}")
        if current_answer is None:
            finding = run_data.get(item['question'])
            if finding: current_answer = finding.get('answer')
        if current_answer in counts:
            counts[current_answer] += 1
    return counts

def create_donut_chart(counts):
    labels = list(counts.keys())
    values = list(counts.values())
    colors = ['#2ca02c', '#d62728', '#ff7f0e']
    fig = go.Figure(data=[go.Pie(labels=labels, values=values, hole=.4, marker_colors=colors, textinfo='value+percent', hoverinfo='label+percent')])
    fig.update_layout(showlegend=True, legend=dict(orientation="h", yanchor="bottom", y=-0.2, xanchor="center", x=0.5), margin=dict(l=20, r=20, t=20, b=20), height=250)
    return fig

# --- WORD REPORT GENERATION ---
def generate_word_report(run_id: str, findings: list) -> io.BytesIO:
    document = Document()
    document.add_heading(f'Audit Remediation Report for: {run_id}', level=1)
    
    document.add_paragraph(
        f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}. "
        "This report lists items requiring action."
    )
    
    actionable_findings = [f for f in findings if f.get('answer') in ["No", "Partial"]]
    
    if not actionable_findings:
        document.add_paragraph("\n✅ No actionable findings were identified for this audit run.")
    else:
        document.add_heading('Actionable Findings', level=2)
        
        # --- CHANGE: Loop and create paragraphs instead of a table ---
        for finding in actionable_findings:
            # Add the question as a distinct heading
            p = document.add_paragraph()
            p.add_run('Question: ').bold = True
            p.add_run(finding.get('question', ''))
            
            # Add the finding (answer)
            p = document.add_paragraph()
            p.add_run('Finding: ').bold = True
            run = p.add_run(finding.get('answer', ''))
            # Optional: Add color to the finding
            # This is more complex and depends on specific library versions.
            # Sticking to bold for simplicity and reliability.

            # Add the explanation
            p = document.add_paragraph()
            p.add_run('Auditronauts Explanation: ').bold = True
            p.add_run(finding.get('explanation', ''))
            
            # Add a separator
            document.add_paragraph("---")

    # Save to a byte stream
    doc_io = io.BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io